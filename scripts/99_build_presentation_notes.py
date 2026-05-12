"""
Builds a printable quick-reference card (Word .docx) for the live
presentation.  Saves to docs/Presentation_Notes.docx.

Layout intent:
* Each slide gets a compact block: ON SLIDE / SAY / IF ASKED / TIMING.
* Bold labels so the eye can lock onto any slide in under 2 seconds.
* No long paragraphs — every line is a short bullet you can read aloud.
"""

import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT_PATH = os.path.join(
    os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
    "docs", "Presentation_Notes.docx",
)

# Cyberdeck-ish palette in Word
CYAN   = RGBColor(0x2A, 0x8A, 0x9A)
RED    = RGBColor(0xB0, 0x40, 0x46)
GOLD   = RGBColor(0x8A, 0x70, 0x50)
GREY   = RGBColor(0x55, 0x55, 0x55)
DARK   = RGBColor(0x18, 0x18, 0x18)

SLIDES = [
    {
        "n": 1,
        "title": "TITLE",
        "on_slide": "Project name, your name, course code, date.",
        "say": [
            "Hi, I'm presenting USCF Cyberdeck Scout — a chess opponent scouting and game-outcome prediction project.",
            "I'll walk through the data, methods, results, and then do a live demo of the dashboard.",
        ],
        "asked": [],
        "timing": "~20 sec",
    },
    {
        "n": 2,
        "title": "THE PROBLEM",
        "on_slide": "Why look beyond rating. Online chess boom. Modern underrated problem.",
        "say": [
            "Most players prep for an opponent by checking one number — their rating.",
            "Rating doesn't show recent form, activity, upset history, or momentum.",
            "Since 2020, many players improve online much faster than their official rating catches up.",
            "So a 'modest-rated' opponent may actually play much stronger than the number suggests.",
            "This project builds a real scouting picture, not just a rating glance.",
        ],
        "asked": [],
        "timing": "~40 sec",
    },
    {
        "n": 3,
        "title": "THE DATASET  ·  NOVEL, NOT KAGGLE",
        "on_slide": "USCF MSA crosstables. Custom scraper. 13,305 raw → 8,514 clean. 179 players · 3,020 events.",
        "say": [
            "Our dataset is novel — it's not from Kaggle.",
            "We scraped public United States Chess Federation tournament crosstables ourselves with a custom Python parser.",
            "After cleaning we have 8,514 modeled games, 179 focal players, 3,020 tournaments, spanning 2001–2025.",
            "Target is binary — did the focal player win — with classes balanced at 49 to 51.",
        ],
        "asked": [
            "Why not Lichess or Chess.com? — Those are online play. We wanted official over-the-board USCF results.",
            "How long did scraping take? — Several batches with a 2-second polite delay; all cached locally so the deployed app never re-scrapes.",
        ],
        "timing": "~30 sec",
    },
    {
        "n": 4,
        "title": "THE PARSER BUG  ·  DATA-QUALITY SAGA",
        "on_slide": "Red warning panel. Symptom → root cause → fix. Before/after numbers.",
        "say": [
            "This is the most important slide of the deck — it's where domain knowledge saved the project.",
            "During EDA I noticed rating_diff correlated NEGATIVELY with winning. That's impossible in chess — higher-rated players should win more, not less.",
            "I traced it to a regex in the parser. It was unanchored — it grabbed the first 4 digits of the 8-digit USCF ID instead of the rating after the 'R:' token.",
            "Example: USCF ID 12640800 was being stored as 'rating 1264' — that's the player's ID prefix, not their skill.",
            "Fix: anchor the regex on 'R:'. Re-parsed all 4,292 cached HTML pages. No re-scrape needed.",
            "Before / After: rating_diff correlation went from −0.13 to +0.49. Elo baseline AUC jumped from 0.37 (worse than random) to 0.825. Random Forest from 0.67 to 0.82.",
            "Takeaway: model accuracy alone wouldn't have caught this. Domain knowledge plus EDA correlations did.",
        ],
        "asked": [
            "How did you know it was the regex? — Spot-checks showed 'ratings' that exactly matched the first 4 digits of opponent IDs. Pattern-matched and confirmed in source.",
            "Did this delay the project? — Re-parsing the cached HTML took ~10 minutes. No re-scraping. Documented in the repo's DATA_FIX_NOTE.",
            "How did you reset the pipeline? — Wrote a re-parse-from-cache script that walks all 3,020 crosstables and rebuilds the raw CSVs in place.",
        ],
        "timing": "~75 sec — SLOW DOWN HERE. This is your differentiator.",
    },
    {
        "n": 5,
        "title": "EDA  ·  WHAT THE DATA LOOKS LIKE",
        "on_slide": "W/D/L 48.8/19.9/31.3 · rating range 227–2861 · rating_diff histogram · redundancy note.",
        "say": [
            "Classes are balanced — 49% wins, 51% not-wins — so accuracy is actually a useful metric here.",
            "Rating range is 227 to 2861 — that's kids playing in scholastic events up to grandmasters.",
            "The rating-difference histogram is a single bell curve centered near zero — Swiss-system tournaments pair similar-rated opponents.",
            "Redundancy audit: rating_diff equals player_rating minus opponent_rating exactly. Using all three in a linear model creates perfect multicollinearity, so the benchmark uses rating_diff alone.",
        ],
        "asked": [
            "Why unimodal at zero? — Tournament directors pair similar-rated players in Swiss systems.",
            "Why exclude raw ratings from the benchmark? — Tree models can use all three because splits don't suffer from collinearity. For the logistic benchmark I keep just rating_diff.",
        ],
        "timing": "~45 sec",
    },
    {
        "n": 6,
        "title": "LITERATURE REVIEW",
        "on_slide": "Elo · Glicko · Glicko-2 · TrueSkill · DeepChess · Maharaj — plus 15 refs in the repo.",
        "say": [
            "We grounded the project in prior chess prediction work.",
            "Statistical rating systems: Elo from 1978, Glicko and Glicko-2 by Glickman, and TrueSkill from Microsoft Research.",
            "Chess-specific ML: DeepChess in 2016, plus Levitt, Stefani, and Maharaj — all of which agree that rating dominates and ML matches Elo within roughly 0.01 AUC on tabular pre-game features.",
            "Full 15-reference review is in the repo.",
        ],
        "asked": [
            "What did Glicko add over Elo? — Each player's rating is a distribution with uncertainty; activity shrinks it, inactivity grows it. We use activity counts as a proxy.",
        ],
        "timing": "~25 sec — keep it brief.",
    },
    {
        "n": 7,
        "title": "METHODOLOGY  ·  HOW WE TRAINED",
        "on_slide": "Chronological 70/15/15 · TimeSeriesSplit(5) · GridSearchCV · 4 models compared.",
        "say": [
            "Splits are chronological — never random — to prevent leaking future games into past predictions.",
            "70% train, 15% validation, 15% test. The test set is 1,278 unseen games from late 2024 through October 2025.",
            "Inside the training set I use TimeSeriesSplit with 5 folds — each fold's validation block sits strictly after its training block.",
            "Hyperparameters tuned with GridSearchCV on mean CV AUC.",
            "Four approaches compared: Logistic Regression as benchmark, Random Forest, Gradient Boosting, plus the closed-form Elo formula.",
            "Metrics reported on all three splits: AUC, F1, Accuracy, Precision, Recall.",
        ],
        "asked": [
            "Why time-aware CV instead of regular k-fold? — Random k-fold lets the model train on 2024 and test on 2010. That's leakage. Time-aware fixes it.",
            "What hyperparameters were tuned? — RF: n_estimators, max_depth, min_samples_split, min_samples_leaf. GB: n_estimators, learning_rate, max_depth.",
        ],
        "timing": "~50 sec",
    },
    {
        "n": 8,
        "title": "MODEL RESULTS",
        "on_slide": "Table: Logistic 0.8246 · RF 0.8234 · GB 0.8248 · Elo 0.8250 — all test AUC.",
        "say": [
            "Test set is 1,278 games the model never saw during training or tuning.",
            "All four approaches land within a tiny window — Test AUC between 0.82 and 0.83.",
            "Gradient Boosting is the best ML model at 0.8248. Elo baseline ties at 0.8250.",
            "Per-fold AUC standard deviation is under 0.015 across the 5 folds — these results are stable, not lucky.",
        ],
        "asked": [
            "Are the differences statistically significant? — No. Per-fold std-dev is 0.015, larger than the inter-model gap. I treat all four as tied.",
            "Why bother with ML if Elo ties? — Raw prediction, no. But the ML pipeline carries recent activity and form features that power the scouting layer.",
            "What metric matters most? — AUC, because the classes are balanced and we want to rank matchups by win probability.",
        ],
        "timing": "~60 sec — most important slide.",
    },
    {
        "n": 9,
        "title": "HONEST FINDING  ·  ML MATCHES ELO",
        "on_slide": "ML matches Elo. Rating dominates ~55% of importance. Recency adds ~13%.",
        "say": [
            "The honest finding: machine learning matches Elo on this task.",
            "That's not a failure — it's what the chess literature predicts. Rating already captures most of what's predictable.",
            "Feature importance breakdown: rating contributes about 55%, recent activity and form together about 13%.",
            "Story: rating dominates, momentum fine-tunes.",
            "The product value lives one layer up — in the scouting report I'll demo next.",
        ],
        "asked": [],
        "timing": "~40 sec — PAUSE after 'ML matches Elo'. The honesty earns credibility.",
    },
    {
        "n": 10,
        "title": "EXTRA CREDIT  ·  CLAUDE LLM END-TO-END",
        "on_slide": "Claude Opus 4.7 · 60-row sample · AUC 0.7239 · F1 0.7463 (highest of any).",
        "say": [
            "For extra credit I ran a generative AI system end-to-end — Claude Opus 4.7.",
            "Same held-out test set, 60-row sample with a fixed seed for reproducibility.",
            "Claude reasoned through each row: Elo base plus form and activity adjustments. Every reasoning trace is saved alongside the predictions.",
            "Result: lower AUC than the ML pipeline at 0.72, but the highest F1 of any approach at 0.7463. Recall 0.93.",
            "Interpretation: the LLM is aggressive about predicting wins. Strong on threshold decisions, weaker on probability calibration than supervised ML.",
        ],
        "asked": [
            "Why only 60 rows? — Per-row reasoning is expensive at inference time. Sixty is enough for a meaningful sample — standard error on AUC at n=60 is about ±0.06.",
            "Did Claude cheat? — Every reasoning trace is in claude_inline_predictions.csv in the repo. You can audit them line by line.",
            "Why higher recall? — When the LLM sees a positive rating advantage plus any recent activity, it tends to round up to 'win.' Supervised ML is more conservative.",
        ],
        "timing": "~60 sec",
    },
    {
        "n": 11,
        "title": "PRODUCT LAYER  ·  SEGUE TO DEMO",
        "on_slide": "0–100 Underrated Potential, 6 subscores, scouting report features.",
        "say": [
            "Built on top of the model: a 0-to-100 Underrated Potential score per player.",
            "Six interpretable subscores — Upset, Recent Form, Momentum, Schedule, Activity, Volatility.",
            "Plus travel and geography, time-control profile, and auto-generated 'why this player might be dangerous' bullets.",
            "Plus practical chess prep tips per matchup.",
            "Let me show you the live dashboard now.",
        ],
        "asked": [],
        "timing": "~30 sec — then JUMP TO DEMO.",
    },
    {
        "n": 12,
        "title": "LIMITS & FUTURE WORK",
        "on_slide": "Disclaimers. Travel approximation. Future: live USCF lookup, supervised underrated model.",
        "say": [
            "Important disclaimers: the Underrated Potential score is a scouting signal, not an accusation. We never say 'smurf.'",
            "Rating remains the strongest single predictor. We add context, not replacement.",
            "Travel distances are approximate — USCF doesn't expose event GPS, so we geocode at city or state level.",
            "Future work: live USCF lookup on demand, deeper geocoding, and a supervised 'predict rating gain' model once we parse post-game ratings.",
        ],
        "asked": [],
        "timing": "~30 sec",
    },
    {
        "n": 13,
        "title": "RUBRIC MAP + TAKE-HOME",
        "on_slide": "Rubric checklist + the one-line quote.",
        "say": [
            "Quick rubric check: novel dataset from USCF, not Kaggle. Cleaning with the parser repair. EDA with redundancy audit. Logistic Regression benchmark. Random Forest and Gradient Boosting tuned with time-aware k-fold cross-validation. Elo plus Claude for extra credit.",
            "Take-home line: rating tells you who is favored — scouting tells you what kind of opponent you're actually facing.",
            "Repo and live app links are on the slide. Happy to take questions.",
        ],
        "asked": [],
        "timing": "~30 sec",
    },
]


def add_para(doc, text, bold=False, color=None, size=11, before=0, after=2):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color
    return p


def add_bullet(doc, text, indent=0.25, size=10.5, color=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(indent)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    run = p.runs[0] if p.runs else p.add_run("")
    if not p.runs:
        run = p.add_run(text)
    else:
        run.text = text
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color


def build():
    doc = Document()

    # Tighter margins so more fits on a printed page
    for section in doc.sections:
        section.top_margin = Inches(0.55)
        section.bottom_margin = Inches(0.55)
        section.left_margin = Inches(0.65)
        section.right_margin = Inches(0.65)

    # ---- Title page (compact, single page header) ----
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = t.add_run("PRESENTATION NOTES")
    r.bold = True; r.font.size = Pt(22); r.font.color.rgb = CYAN

    sub = doc.add_paragraph()
    sr = sub.add_run("USCF Cyberdeck Scout  ·  ML Final Project  ·  Quick-Reference Card")
    sr.font.size = Pt(11); sr.font.color.rgb = GREY

    intro = doc.add_paragraph()
    ir = intro.add_run(
        "Print and keep beside you. Each slide block is structured the same way: "
        "what's on the slide, what to say, what to answer if asked. "
        "Bold labels let you find any slide in under two seconds."
    )
    ir.font.size = Pt(10); ir.font.color.rgb = DARK

    # Quick legend
    leg = doc.add_paragraph()
    leg_r = leg.add_run("LEGEND  ")
    leg_r.bold = True; leg_r.font.size = Pt(9); leg_r.font.color.rgb = CYAN
    leg_r2 = leg.add_run(
        "ON SLIDE = what the audience sees   ·   "
        "SAY = your spoken line   ·   "
        "IF ASKED = likely Q&A   ·   "
        "TIMING = pace target"
    )
    leg_r2.font.size = Pt(9); leg_r2.font.color.rgb = GREY

    # Separator
    sep = doc.add_paragraph()
    sep_r = sep.add_run("─" * 90)
    sep_r.font.size = Pt(8); sep_r.font.color.rgb = CYAN
    sep.paragraph_format.space_after = Pt(4)

    # ---- One block per slide ----
    for s in SLIDES:
        # Slide header bar
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(10)
        h.paragraph_format.space_after = Pt(2)
        num_run = h.add_run(f"SLIDE {s['n']:02d}")
        num_run.bold = True; num_run.font.size = Pt(13); num_run.font.color.rgb = CYAN
        spacer = h.add_run("   ·   ")
        spacer.font.size = Pt(13); spacer.font.color.rgb = GREY
        title_run = h.add_run(s["title"])
        title_run.bold = True; title_run.font.size = Pt(13); title_run.font.color.rgb = DARK

        # ON SLIDE
        on = doc.add_paragraph()
        on.paragraph_format.space_after = Pt(2)
        on_label = on.add_run("ON SLIDE   ")
        on_label.bold = True; on_label.font.size = Pt(9.5); on_label.font.color.rgb = GOLD
        on_body = on.add_run(s["on_slide"])
        on_body.font.size = Pt(10); on_body.font.color.rgb = DARK

        # SAY
        say_p = doc.add_paragraph()
        say_p.paragraph_format.space_after = Pt(2)
        say_label = say_p.add_run("SAY")
        say_label.bold = True; say_label.font.size = Pt(9.5); say_label.font.color.rgb = CYAN
        for line in s["say"]:
            b = doc.add_paragraph(style="List Bullet")
            b.paragraph_format.left_indent = Inches(0.25)
            b.paragraph_format.space_before = Pt(0)
            b.paragraph_format.space_after = Pt(1)
            br = b.runs[0] if b.runs else b.add_run("")
            if not b.runs:
                br = b.add_run(line)
            else:
                br.text = line
            br.font.size = Pt(10.5)
            br.font.color.rgb = DARK

        # IF ASKED
        if s["asked"]:
            ifp = doc.add_paragraph()
            ifp.paragraph_format.space_before = Pt(2)
            ifp.paragraph_format.space_after = Pt(2)
            if_label = ifp.add_run("IF ASKED")
            if_label.bold = True; if_label.font.size = Pt(9.5); if_label.font.color.rgb = RED
            for line in s["asked"]:
                b = doc.add_paragraph(style="List Bullet")
                b.paragraph_format.left_indent = Inches(0.25)
                b.paragraph_format.space_before = Pt(0)
                b.paragraph_format.space_after = Pt(1)
                br = b.runs[0] if b.runs else b.add_run("")
                if not b.runs:
                    br = b.add_run(line)
                else:
                    br.text = line
                br.font.size = Pt(9.5)
                br.font.color.rgb = GREY

        # TIMING
        tm = doc.add_paragraph()
        tm.paragraph_format.space_before = Pt(2)
        tm.paragraph_format.space_after = Pt(4)
        tm_label = tm.add_run("TIMING   ")
        tm_label.bold = True; tm_label.font.size = Pt(9); tm_label.font.color.rgb = GOLD
        tm_body = tm.add_run(s["timing"])
        tm_body.font.size = Pt(9.5); tm_body.font.color.rgb = DARK
        tm_body.italic = True

    # ---- Footer crib: the three core lines ----
    sep2 = doc.add_paragraph()
    sep2_r = sep2.add_run("─" * 90)
    sep2_r.font.size = Pt(8); sep2_r.font.color.rgb = CYAN
    sep2.paragraph_format.space_before = Pt(8)

    crib_h = doc.add_paragraph()
    crib_r = crib_h.add_run("EMERGENCY CRIB  ·  IF YOU FORGET EVERYTHING ELSE, SAY THESE THREE")
    crib_r.bold = True; crib_r.font.size = Pt(11); crib_r.font.color.rgb = RED

    for i, line in enumerate([
        "Dataset is NOVEL — scraped from USCF, not Kaggle. (Slide 3)",
        "EDA + chess theory caught a parser bug that model accuracy would have missed. (Slide 4)",
        "ML matches Elo at AUC 0.82. Rating dominates. Product value is the scouting layer. (Slide 9)",
    ], start=1):
        b = doc.add_paragraph()
        b.paragraph_format.left_indent = Inches(0.2)
        b.paragraph_format.space_after = Pt(2)
        num = b.add_run(f"{i}.  ")
        num.bold = True; num.font.size = Pt(11); num.font.color.rgb = CYAN
        body = b.add_run(line)
        body.font.size = Pt(11); body.font.color.rgb = DARK

    os.makedirs(os.path.dirname(OUT_PATH), exist_ok=True)
    doc.save(OUT_PATH)
    print(f"Saved: {OUT_PATH}")


if __name__ == "__main__":
    build()
